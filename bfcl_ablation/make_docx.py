#!/usr/bin/env python3
"""Render each experiment report to a standalone .docx.

The Markdown in `experiments/` is the source of truth for the narrative. This adds
what a circulated document needs and a repo file does not: a cover block naming the
pack, model and metric contract the numbers were produced under, a full data appendix
read back from `results/<ARM>/metrics.json` rather than retyped, and a reproduction
section with artifact digests.

Retyping numbers into a report is how a report drifts from its data, so every figure
in the appendix is read from the JSON at render time.

    PYTHONPATH=src:. python3 bfcl_ablation/make_docx.py
    PYTHONPATH=src:. python3 bfcl_ablation/make_docx.py --arm a4
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from pathlib import Path
from typing import Any

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402

from bfcl_ablation import common  # noqa: E402

EXPERIMENTS = common.ABLATION_ROOT / "experiments"
DOCX_OUT = common.ABLATION_ROOT / "reports"

ARMS = {
    "a0": ("A0", "Human baseline", "Measurement stage on the unmodified pack"),
    "a1": ("A1", "Deterministic simplification", "How much friction comes off with no model"),
    "a2": ("A2", "LLM surface generation", "Wording only; ground truth frozen"),
    "a3": ("A3", "LLM task generation", "Semantics; controlled policy sampler"),
    "a4": ("A4", "LLM assertions and the mutation gate", "Are the assertions checking anything"),
    "a5": ("A5", "Target-model evaluation across wordings", "Does a benchmark conclusion survive a paraphrase"),
    "a6": ("A6", "Backend mutation gate", "Is the oracle itself falsifiable"),
    "a2_rerun": ("A2_rerun", "The reproduction run", "Does A2 reproduce, and what does a passing run hide"),
    "findings": ("ALL", "Findings across the ladder", "Cross-arm synthesis"),
}

MONO = "Consolas"
CODE_GREY = RGBColor(0x24, 0x29, 0x2E)
MUTED = RGBColor(0x60, 0x6A, 0x76)


# --------------------------------------------------------------------------------
# inline markdown
# --------------------------------------------------------------------------------

_INLINE = re.compile(
    r"(\*\*.+?\*\*)"      # bold
    r"|(`[^`]+?`)"        # code
    r"|(\*[^*]+?\*)"      # italic
    r"|(\[[^\]]+?\]\([^)]+?\))"  # link
)


def add_inline(paragraph, text: str) -> None:
    """Write text into a paragraph, honouring bold, italic, code and links.

    A link renders as its label followed by the target in parentheses only when the
    target is a URL. A repo-relative path is dropped: it is noise on paper and the
    reproduction section already lists every artifact path.
    """
    for chunk in filter(None, _INLINE.split(text)):
        if chunk.startswith("**") and chunk.endswith("**"):
            paragraph.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = MONO
            run.font.size = Pt(9.5)
            run.font.color.rgb = CODE_GREY
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            paragraph.add_run(chunk[1:-1]).italic = True
        elif chunk.startswith("[") and "](" in chunk:
            label, _, target = chunk[1:-1].partition("](")
            add_inline(paragraph, label)
            if target.startswith("http"):
                run = paragraph.add_run(f" ({target})")
                run.font.size = Pt(8.5)
                run.font.color.rgb = MUTED
        else:
            paragraph.add_run(chunk)


# --------------------------------------------------------------------------------
# block markdown
# --------------------------------------------------------------------------------


def _is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def _cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _alignments(separator: str) -> list[int]:
    result = []
    for cell in _cells(separator):
        if cell.endswith(":") and cell.startswith(":"):
            result.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif cell.endswith(":"):
            result.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            result.append(WD_ALIGN_PARAGRAPH.LEFT)
    return result


def add_table(document: Document, rows: list[list[str]], alignments: list[int]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            if column_index >= len(table.columns):
                continue
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = alignments[column_index] if column_index < len(alignments) else None
            add_inline(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if row_index == 0:
                    run.bold = True
    document.add_paragraph()


def render_markdown(document: Document, text: str, *, skip_first_h1: bool = True) -> None:
    lines = text.splitlines()
    index = 0
    seen_h1 = False

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            index += 1
            body: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                body.append(lines[index])
                index += 1
            index += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(18)
            paragraph.paragraph_format.space_after = Pt(10)
            run = paragraph.add_run("\n".join(body))
            run.font.name = MONO
            run.font.size = Pt(8.5)
            run.font.color.rgb = CODE_GREY
            continue

        if _is_table_row(stripped) and index + 1 < len(lines) and set(lines[index + 1].strip()) <= set("|-: "):
            header = _cells(stripped)
            alignments = _alignments(lines[index + 1])
            index += 2
            rows = [header]
            while index < len(lines) and _is_table_row(lines[index].strip()):
                rows.append(_cells(lines[index].strip()))
                index += 1
            width = len(header)
            rows = [row[:width] + [""] * (width - len(row)) for row in rows]
            add_table(document, rows, alignments)
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            if level == 1 and not seen_h1:
                seen_h1 = True
                if skip_first_h1:
                    index += 1
                    continue
            heading = document.add_heading("", level=min(level, 4))
            add_inline(heading, title)
            index += 1
            continue

        if stripped in {"---", "***", "___"}:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(2)
            run = paragraph.add_run("_" * 78)
            run.font.color.rgb = MUTED
            run.font.size = Pt(7)
            index += 1
            continue

        if stripped.startswith("> "):
            block = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                block.append(lines[index].strip().lstrip(">").strip())
                index += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(24)
            add_inline(paragraph, " ".join(block))
            for run in paragraph.runs:
                run.italic = True
            continue

        if re.match(r"^[-*] ", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            index += 1
            continue

        if re.match(r"^\d+\. ", stripped):
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, stripped.split(". ", 1)[1])
            index += 1
            continue

        block = []
        while index < len(lines) and lines[index].strip() and not re.match(
            r"^(#|\||```|>|[-*] |\d+\. |---$)", lines[index].strip()
        ):
            block.append(lines[index].strip())
            index += 1
        if block:
            add_inline(document.add_paragraph(), " ".join(block))


# --------------------------------------------------------------------------------
# appendices, read from the JSON rather than retyped
# --------------------------------------------------------------------------------


def simple_table(document: Document, header: list[str], rows: list[list[Any]]) -> None:
    if not rows:
        document.add_paragraph("(no rows)")
        return
    alignments = [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.RIGHT] * (len(header) - 1)
    add_table(document, [header] + [[str(c) for c in row] for row in rows], alignments)


def appendix_a0(
    document: Document,
    metrics: dict[str, Any],
    *,
    prefix: str = "A.",
    level: int = 3,
    include_sweep: bool = True,
) -> None:
    """Render the shared benchmark-property tables.

    A3 produces the same metric schema, so it reuses this block rather than
    duplicating it — `prefix` keeps the nested numbering distinct, and `include_sweep`
    is off there because the budget sweep is an A0 experiment.
    """
    friction = metrics["authoring_friction"]
    document.add_heading(f"{prefix}1 Authoring cost by file", level=level)
    simple_table(
        document,
        ["file", "lines"],
        [[name, count] for name, count in friction["loc_by_file"].items()],
    )

    distribution = metrics["distribution"]
    document.add_heading(f"{prefix}2 Joint (category x policy) matrix", level=level)
    document.add_paragraph(
        "Cell values are task counts. 'unwritten' means feasible but not authored; "
        "'structural' means no tool in the category's universe can satisfy the policy."
    )
    policies = distribution["policies"]
    rows = []
    for category, row in distribution["matrix"].items():
        line = [category]
        for policy in policies:
            cell = row[policy]
            status = cell.get("status")
            line.append(
                str(cell["tasks"]) if not status
                else ("--" if status == "empty_structural" else ".")
            )
        rows.append(line)
    simple_table(document, ["category"] + policies, rows)

    document.add_heading(f"{prefix}3 Policy distribution", level=level)
    simple_table(
        document,
        ["turn_policy", "tasks", "share"],
        [
            [policy, count, f"{distribution['policy_task_share'][policy]:.1%}"]
            for policy, count in distribution["policy_task_counts"].items()
        ],
    )

    coverage = metrics["coverage"]
    document.add_heading(f"{prefix}4 Fixture coverage", level=level)
    simple_table(
        document,
        ["collection", "primary key", "rows", "bound", "coverage", "never bound"],
        [
            [
                name,
                entry["primary_key"],
                entry["rows"],
                entry["entities_bound"],
                f"{entry['coverage']:.0%}" if entry["coverage"] is not None else "n/a",
                ", ".join(entry["entities_never_bound"][:8]) or "-",
            ]
            for name, entry in coverage["fixtures"].items()
        ],
    )
    document.add_paragraph(
        f"Tool call counts: {json.dumps(coverage['tool_call_counts'])}. "
        f"Tools never called: {', '.join(coverage['tools_never_called']) or 'none'}."
    )

    surface = metrics["surface"]
    document.add_heading(f"{prefix}5 Surface diversity by category", level=level)
    simple_table(
        document,
        ["category", "tasks", "distinct masked", "ratio"],
        [
            [name, entry["tasks"], entry["distinct_masked"],
             f"{entry['ratio']:.2f}" if entry["ratio"] is not None else "n/a"]
            for name, entry in surface["by_category"].items()
        ],
    )
    document.add_heading(f"{prefix}6 Most repeated masked utterances", level=level)
    simple_table(
        document,
        ["tasks", "utterance"],
        [[entry["tasks"], entry["utterance"]] for entry in surface["most_repeated_masked"]],
    )

    funnel = metrics["funnel"]
    document.add_heading(f"{prefix}7 Publish funnel", level=level)
    simple_table(
        document,
        ["stage", "rows", "survival", "lost here", "drop reasons"],
        [
            [
                step["stage"],
                step["rows"],
                f"{step['survival_from_expand']:.1%}",
                step["lost_here"],
                json.dumps(step["drop_reasons"]) if step["drop_reasons"] else "-",
            ]
            for step in funnel["steps"]
        ],
    )

    sweep = common.read_json(common.RESULTS / "budget_sweep.json") if include_sweep else None
    if sweep:
        document.add_heading(f"{prefix}8 Budget sweep", level=level)
        document.add_paragraph(
            "The only generation knob available. It moves entity coverage and leaves "
            "surface diversity unchanged."
        )
        simple_table(
            document,
            ["tasks_per_category", "tasks", "published", "entities bound", "distinct surfaces", "single_turn share"],
            [
                [
                    row["budget"], row["tasks"], row["published"],
                    f"{row['entities_bound']}/{row['entities_total']}",
                    row["distinct_surfaces"], f"{row['single_turn_share']:.1%}",
                ]
                for row in sweep
            ],
        )


def appendix_a1(document: Document, metrics: dict[str, Any]) -> None:
    simplification = metrics.get("simplification") or {}
    baseline_loc = simplification.get("loc_baseline") or {}
    arm_loc = simplification.get("loc_a1") or {}
    document.add_heading("A.1 Authoring cost, A0 against A1", level=3)
    simple_table(
        document,
        ["file", "A0", "A1", "saved"],
        [[name, before, arm_loc.get(name, 0), before - arm_loc.get(name, 0)]
         for name, before in baseline_loc.items()],
    )

    findings = (simplification.get("shrink") or {}).get("template_findings", [])
    manifest_findings = (simplification.get("shrink") or {}).get("manifest_findings", [])

    # The 46-row verdict list below answers "which field on which template"; this
    # rollup answers "which kinds of field came off at all", which is the question a
    # reader actually has first.
    rollup: dict[str, dict[str, int]] = {}
    for finding in findings + manifest_findings:
        rollup.setdefault(finding["field"], {}).setdefault(finding["verdict"], 0)
        rollup[finding["field"]][finding["verdict"]] += 1
    document.add_heading("A.2 Shrink summary by field", level=3)
    verdicts = ["derived_exact", "surface_changed", "not_derivable", "compiler_error"]
    present = [v for v in verdicts if any(v in counts for counts in rollup.values())]
    simple_table(
        document,
        ["field"] + present,
        [
            [field] + [rollup[field].get(v, 0) for v in present]
            for field in sorted(rollup, key=lambda f: -sum(rollup[f].values()))
        ],
    )

    document.add_heading("A.3 Field-level derivability verdicts", level=3)
    simple_table(
        document,
        ["where", "field", "verdict", "detail"],
        [
            [f.get("template_id", "manifest"), f["field"], f["verdict"], (f.get("detail") or "")[:120]]
            for f in findings + manifest_findings
        ],
    )

    equivalence = metrics.get("equivalence") or {}
    document.add_heading("A.4 Equivalence checks", level=3)
    simple_table(
        document,
        ["check", "result"],
        [
            ["set(task_id) equal", equivalence.get("task_ids", {}).get("equal")],
            ["expected_tool_calls equal", equivalence.get("expected_tool_calls", {}).get("equal")],
            ["conversation_plans equal", equivalence.get("conversation_plans", {}).get("equal")],
            ["validation-case coverage held", equivalence.get("validation_case_coverage", {}).get("held")],
            ["opening user turn identical", equivalence.get("surface", {}).get("first_turns_identical")],
            ["gold eligibility preserved", equivalence.get("publication", {}).get("gold_preserved")],
            ["VERDICT", equivalence.get("verdict")],
        ],
    )

    config = metrics.get("config_minimization") or {}
    if config:
        document.add_heading("A.5 Run-config minimization", level=3)
        document.add_paragraph(
            f"Run config {config.get('config_lines_before')} -> {config.get('config_lines_after')} lines. "
            f"Verdict against the baseline: {config.get('equivalence', {}).get('verdict')}."
        )
        simple_table(
            document,
            ["setting dropped"],
            [[name] for name in (config.get("minimization") or {}).get("dropped", [])],
        )


def appendix_a4(document: Document, metrics: dict[str, Any]) -> None:
    arms = metrics.get("arms") or {}
    document.add_heading("A.1 False acceptance by operator class (strict operators only)", level=3)
    classes = ["call_level", "argument_level", "state_level"]
    rows = []
    for name, arm in arms.items():
        line = [name]
        for cls in classes:
            entry = (arm.get("by_class_strict") or {}).get(cls, {})
            line.append(
                f"{entry.get('false_acceptance_rate', float('nan')):.3f} "
                f"({entry.get('false_accept')}/{entry.get('trials')})"
            )
        gold = arm.get("gold") or {}
        line.append(f"{gold.get('passed')}/{gold.get('instances')}")
        rows.append(line)
    simple_table(document, ["assertions"] + classes + ["unmutated pass"], rows)

    document.add_heading("A.2 False acceptance by operator", level=3)
    operators = sorted((arms.get("human") or {}).get("by_operator") or {})
    rows = []
    for operator in operators:
        line = [operator]
        for name in arms:
            entry = (arms[name].get("by_operator") or {}).get(operator, {})
            rate = entry.get("false_acceptance_rate")
            line.append(f"{rate:.3f} ({entry.get('trials', 0)})" if rate is not None else "-")
        rows.append(line)
    simple_table(document, ["operator"] + list(arms), rows)

    document.add_heading("A.3 False acceptance by assertion (human suite)", level=3)
    by_assertion = (arms.get("human") or {}).get("by_assertion") or {}
    simple_table(
        document,
        ["assertion", "trials", "detected", "false accept", "FA rate"],
        [
            [name, e.get("trials"), e.get("detected"), e.get("false_accept"),
             f"{e.get('false_acceptance_rate', 0):.3f}"]
            for name, e in sorted(by_assertion.items())
        ],
    )

    document.add_heading("A.4 Operator inventory", level=3)
    simple_table(
        document,
        ["operator", "class", "delivery mode", "mutations", "tasks"],
        [
            [name, e.get("op_class"), e.get("mode"), e.get("mutations"), e.get("tasks")]
            for name, e in sorted((metrics.get("operator_inventory") or {}).items())
        ],
    )
    document.add_paragraph(
        f"Assertion suite sizes (lines): human {metrics.get('assertions_loc', {}).get('human')}; "
        + "; ".join(
            f"{k} {v.get('assertions_loc')}"
            for k, v in (metrics.get("authoring") or {}).items()
            if isinstance(v, dict) and v.get("assertions_loc")
        )
    )


def appendix_a2(document: Document, metrics: dict[str, Any]) -> None:
    design = metrics.get("design") or {}
    document.add_heading("A.1 Design", level=3)
    simple_table(
        document,
        ["setting", "value"],
        [
            ["budgets swept", design.get("budgets")],
            ["paraphrase counts (N)", design.get("n_rungs")],
            ["language", design.get("language")],
            ["variant 0", "the authored sentence" if design.get("variant_0_is_authored_sentence") else "generated"],
            ["variant assignment", design.get("assembly")],
        ],
    )

    generation = metrics.get("generation") or {}
    document.add_heading("A.2 Paraphrase generation", level=3)
    simple_table(
        document,
        ["measure", "value"],
        [
            ["requested per template", generation.get("requested_per_template")],
            ["rejection rate", generation.get("rejection_rate")],
            ["rejection reasons", json.dumps(generation.get("rejection_reasons") or {})],
            ["templates with no pool", ", ".join(generation.get("templates_without_pool") or []) or "none"],
        ],
    )

    document.add_heading("A.3 Per-rung results", level=3)
    document.add_paragraph(
        "Ceiling is the sum over templates of min(N, pool size, tasks for that template). "
        "A rung below its ceiling lost variety to variant collision, not to the model."
    )
    rows = []
    for budget_key in sorted(metrics.get("budgets") or {}, key=lambda k: int(k)):
        block = metrics["budgets"][budget_key]
        for rung in block.get("rungs") or []:
            equivalence = rung.get("equivalence") or {}
            rows.append([
                block.get("budget"),
                rung.get("n"),
                rung.get("tasks"),
                rung.get("distinct_masked_surfaces"),
                rung.get("ceiling"),
                "YES" if (equivalence.get("task_ids") or {}).get("equal") else "NO",
                "YES" if (equivalence.get("expected_tool_calls") or {}).get("equal") else "NO",
                "YES" if (equivalence.get("conversation_plans") or {}).get("equal") else "NO",
                rung.get("published_flagged_surface"),
                rung.get("published_substituted_surface"),
            ])
    simple_table(
        document,
        ["budget", "N", "tasks", "distinct masked", "ceiling",
         "task_id equal", "tool_calls equal", "plans equal", "flagged published", "substituted published"],
        rows,
    )

    check = metrics.get("intent_check") or {}
    document.add_heading("A.4 Intent-preservation check", level=3)
    document.add_paragraph(
        "An independent model call reads the sentence and the tool catalogue only — no "
        "expected answer, no required_tools — and names the tools the request needs. "
        "Disagreement flags the sentence."
    )
    simple_table(
        document,
        ["population", "n", "flagged", "rate"],
        [
            ["injected intent shifts (all should be caught)",
             (check.get("shift_recall") or {}).get("n"),
             (check.get("shift_recall") or {}).get("count"),
             (check.get("shift_recall") or {}).get("rate")],
            ["canonical authored sentences (false-alarm floor)",
             (check.get("canonical_false_alarm") or {}).get("n"),
             (check.get("canonical_false_alarm") or {}).get("count"),
             (check.get("canonical_false_alarm") or {}).get("rate")],
            ["generated paraphrases",
             (check.get("paraphrase_false_alarm") or {}).get("n"),
             (check.get("paraphrase_false_alarm") or {}).get("count"),
             (check.get("paraphrase_false_alarm") or {}).get("rate")],
            ["of those, intent actually substituted",
             (check.get("paraphrase_substitution") or {}).get("n"),
             (check.get("paraphrase_substitution") or {}).get("count"),
             (check.get("paraphrase_substitution") or {}).get("rate")],
        ],
    )
    document.add_paragraph(
        f"Flag kinds on paraphrases: {json.dumps(check.get('paraphrase_flag_kinds') or {})}. "
        f"On injected shifts: {json.dumps(check.get('shift_flag_kinds') or {})}. "
        f"Shifts stopped by the mechanical guards before the checker saw them: "
        f"{json.dumps(check.get('shifts_rejected_by_mechanical_guards') or {})}."
    )

    document.add_heading("A.5 Flag rate by turn policy", level=3)
    document.add_paragraph(
        "The flag rate is policy-shaped, which is a property of the checker rather than of "
        "the paraphrases: it sees one opening turn and cannot infer a multi-step plan."
    )
    simple_table(
        document,
        ["turn policy", "paraphrases", "flagged", "rate", "kinds"],
        [
            [policy, entry.get("n"), entry.get("flagged"), entry.get("rate"),
             json.dumps(entry.get("kinds") or {})]
            for policy, entry in sorted((check.get("paraphrase_by_policy") or {}).items())
        ],
    )
    if check.get("caveat"):
        paragraph = document.add_paragraph()
        add_inline(paragraph, check["caveat"])
        for run in paragraph.runs:
            run.italic = True

    llm = metrics.get("llm") or {}
    document.add_heading("A.6 Model provenance", level=3)
    simple_table(
        document,
        ["field", "value"],
        [["model", llm.get("model")], ["endpoint", llm.get("base_url")],
         ["call stats", json.dumps(llm.get("stats") or {})]],
    )


def appendix_a3(document: Document, metrics: dict[str, Any]) -> None:
    proposal = metrics.get("proposal") or {}
    document.add_heading("A.1 Proposal outcome", level=3)
    simple_table(
        document,
        ["measure", "value"],
        [
            ["proposals requested", proposal.get("proposals_requested")],
            ["proposals returned", proposal.get("proposals_returned")],
            ["accepted", proposal.get("accepted")],
            ["accept rate", proposal.get("accept_rate")],
            ["drop buckets", json.dumps(proposal.get("drop_buckets") or {})],
            ["drop stages", json.dumps(proposal.get("drop_stages") or {})],
        ],
    )

    bias = metrics.get("selection_bias") or {}
    failure = bias.get("failure_bias") or {}
    document.add_heading("A.2 Accept rate by policy — bias that survives a controlled sampler", level=3)
    document.add_paragraph(
        "The sampler proposes policies to a target. What survives the accept/drop gate is a "
        "different distribution, and that difference is the finding."
    )
    simple_table(
        document,
        ["turn policy", "proposed", "accepted", "accept rate", "drop buckets"],
        [
            [policy, e.get("proposed"), e.get("accepted"), e.get("accept_rate"),
             json.dumps(e.get("drop_buckets") or {})]
            for policy, e in sorted((failure.get("by_policy") or {}).items())
        ],
    )
    document.add_paragraph(f"Spread between easiest and hardest policy: {failure.get('accept_rate_spread')}.")

    tools = bias.get("tool_choice") or {}
    document.add_heading("A.3 Tool choice against a conditional-uniform null", level=3)
    observed = tools.get("observed") or {}
    expected = tools.get("expected_uniform_within_category") or {}
    ratio = tools.get("ratio_observed_over_expected") or {}
    simple_table(
        document,
        ["tool", "observed", "expected", "obs/exp"],
        [[name, observed.get(name, 0), expected.get(name), ratio.get(name, "-")] for name in sorted(set(observed) | set(expected))],
    )
    document.add_paragraph(
        f"Pearson statistic: {tools.get('chi_square_vs_conditional_uniform')}. "
        f"Total variation from uniform over all tools: {tools.get('tvd_from_uniform_over_all_tools')}. "
        f"Never required: {', '.join(tools.get('tools_never_required') or []) or 'none'}."
    )

    entities = bias.get("entity_choice") or {}
    document.add_heading("A.4 Entity choice", level=3)
    simple_table(
        document,
        ["collection", "rows", "rows bound", "bindings", "TVD from uniform"],
        [
            [name, e.get("rows"), e.get("rows_bound"), e.get("bindings"), e.get("tvd_from_uniform")]
            for name, e in sorted((entities.get("by_collection") or {}).items())
        ],
    )
    document.add_paragraph(
        f"Entities bound: {entities.get('entities_bound')}/{entities.get('entities_total')} "
        f"({entities.get('coverage')})."
    )

    document.add_heading("A.5 Distance from the human-authored mix", level=3)
    rows = []
    for field, entry in (bias.get("vs_a0") or {}).items():
        rows.append([
            field,
            entry.get("tvd_a3_vs_a0"),
            ", ".join(entry.get("in_a0_only") or []) or "-",
            ", ".join(entry.get("in_a3_only") or []) or "-",
        ])
    simple_table(document, ["field", "TVD A3 vs A0", "only in A0", "only in A3"], rows)

    achieved = metrics.get("coverage_achieved") or {}
    spec = metrics.get("coverage_spec") or {}
    document.add_heading("A.6 Coverage against the spec", level=3)
    simple_table(
        document,
        ["measure", "value"],
        [
            ["cells total", spec.get("cells_total")],
            ["cells feasible", achieved.get("cells_feasible")],
            ["cells declared structurally empty", achieved.get("cells_structural_empty")],
            ["cells covered", achieved.get("cells_covered")],
            ["cells meeting target", achieved.get("cells_target_met")],
            ["coverage rate", achieved.get("coverage_rate")],
        ],
    )
    unmet = [c for c in (achieved.get("cells") or []) if c.get("feasible") and not c.get("met")]
    if unmet:
        document.add_paragraph("Feasible cells the model could not fill:")
        simple_table(
            document,
            ["category", "policy", "target", "achieved"],
            [[c.get("category"), c.get("policy"), c.get("target"), c.get("achieved")] for c in unmet],
        )

    probe = metrics.get("backend_probe") or {}
    document.add_heading("A.7 Backend dependency probe", level=3)
    document.add_paragraph(
        "Producer/consumer pairs read off the backend, used to decide whether a "
        "dependent_call cell is feasible rather than guessing from tool count."
    )
    simple_table(
        document,
        ["producer", "consumer", "parameter", "path"],
        [[e.get("producer"), e.get("consumer"), e.get("parameter"), e.get("path")]
         for e in (probe.get("dependency_edges") or [])],
    )

    document.add_heading("A.8 Resulting benchmark", level=3)
    appendix_a0(document, metrics, prefix="A.8.", level=4, include_sweep=False)


def appendix_all(document: Document, _metrics: dict[str, Any] | None = None) -> None:
    """Headline readouts for all five arms, so the synthesis stands alone."""
    document.add_heading("A.1 Headline readouts by arm", level=3)
    rows = []
    for arm, label in (("a0", "A0"), ("a1", "A1"), ("a3", "A3")):
        metrics = common.read_json(common.RESULTS / label / "metrics.json")
        if not metrics:
            continue
        friction = metrics.get("authoring_friction") or {}
        surface = (metrics.get("surface") or {}).get("overall") or {}
        coverage = metrics.get("coverage") or {}
        funnel = metrics.get("funnel") or {}
        rows.append([
            label,
            friction.get("loc_total"),
            friction.get("template_count"),
            funnel.get("expanded"),
            funnel.get("published"),
            surface.get("distinct_masked"),
            f"{coverage.get('fixture_entities_bound')}/{coverage.get('fixture_entities_total')}",
            f"{funnel.get('publish_rate', 0):.0%}",
        ])
    simple_table(
        document,
        ["arm", "authored LOC", "templates", "tasks", "published", "distinct surfaces", "entities bound", "publish rate"],
        rows,
    )
    document.add_paragraph(
        "A2 holds the A0 task set frozen and varies wording only, so its task and entity "
        "counts are A0's by construction; its own table is A.2."
    )

    a2 = common.read_json(common.RESULTS / "A2" / "metrics.json") or {}
    document.add_heading("A.2 A2 — surface diversity by rung", level=3)
    rows = []
    for budget_key in sorted(a2.get("budgets") or {}, key=lambda k: int(k)):
        block = a2["budgets"][budget_key]
        for rung in block.get("rungs") or []:
            equivalence = rung.get("equivalence") or {}
            rows.append([
                block.get("budget"), rung.get("n"), rung.get("tasks"),
                rung.get("distinct_masked_surfaces"), rung.get("ceiling"),
                equivalence.get("verdict", "-"),
            ])
    simple_table(document, ["budget", "N", "tasks", "distinct masked", "ceiling", "verdict"], rows)

    a4 = common.read_json(common.RESULTS / "A4" / "metrics.json") or {}
    document.add_heading("A.3 A4 — false acceptance by operator class (strict)", level=3)
    classes = ["call_level", "argument_level", "state_level"]
    rows = []
    for name, arm in (a4.get("arms") or {}).items():
        line = [name]
        for cls in classes:
            entry = (arm.get("by_class_strict") or {}).get(cls, {})
            rate = entry.get("false_acceptance_rate")
            line.append(f"{rate:.3f} ({entry.get('false_accept')}/{entry.get('trials')})" if rate is not None else "-")
        gold = arm.get("gold") or {}
        line.append(f"{gold.get('passed')}/{gold.get('instances')}")
        rows.append(line)
    simple_table(document, ["assertions"] + classes + ["unmutated pass"], rows)

    sweep = common.read_json(common.RESULTS / "budget_sweep.json")
    if sweep:
        document.add_heading("A.4 Budget sweep (A0)", level=3)
        simple_table(
            document,
            ["tasks_per_category", "tasks", "entities bound", "distinct surfaces", "single_turn share"],
            [
                [r["budget"], r["tasks"], f"{r['entities_bound']}/{r['entities_total']}",
                 r["distinct_surfaces"], f"{r['single_turn_share']:.1%}"]
                for r in sweep
            ],
        )


def appendix_a5(document: Document, metrics: dict[str, Any]) -> None:
    """The paired table and the per-policy breakdown, which is where the result lives.

    The pooled delta averages 18 `single_turn` tasks that did not move with 5
    `confirmation` tasks that dropped 40%, so the per-policy table is not supporting
    detail — it is the finding.
    """
    paired = metrics.get("paired") or {}
    document.add_heading("A.1 Paired outcome, both verdicts", level=3)
    simple_table(
        document,
        ["verdict", "A0", "A2", "delta", "agreement", "discordant", "McNemar p"],
        [
            [
                key,
                row.get("accuracy_a0"),
                row.get("accuracy_a2"),
                row.get("delta"),
                row.get("paired_agreement"),
                row.get("discordant"),
                row.get("mcnemar_p"),
            ]
            for key, row in paired.items()
        ],
    )

    ast = paired.get("ast_match") or {}
    contingency = ast.get("contingency") or {}
    document.add_heading("A.2 Contingency (ast_match)", level=3)
    simple_table(
        document,
        ["", "A2 correct", "A2 wrong"],
        [
            ["A0 correct", contingency.get("both_correct"), contingency.get("a0_only")],
            ["A0 wrong", contingency.get("a2_only"), contingency.get("neither")],
        ],
    )

    document.add_heading("A.3 Per turn policy (ast_match)", level=3)
    simple_table(
        document,
        ["policy", "n", "A0", "A2", "delta", "flipped down", "flipped up"],
        [
            [name, row["n"], row["accuracy_a0"], row["accuracy_a2"], row["delta"],
             row["flipped_down"], row["flipped_up"]]
            for name, row in ((metrics.get("by_turn_policy") or {}).get("ast_match") or {}).items()
        ],
    )

    disagreement = metrics.get("verdict_disagreement") or {}
    document.add_heading("A.4 Verdict disagreement", level=3)
    simple_table(
        document,
        ["case", "count"],
        [
            ["assertions passed while the calls were wrong", disagreement.get("assertion_passed_ast_failed")],
            ["calls right while assertions failed", disagreement.get("ast_passed_assertion_failed")],
        ],
    )


def appendix_a6(document: Document, metrics: dict[str, Any]) -> None:
    """The layer table and the operator table, which carry the whole result.

    The raw survival count is deliberately NOT the headline here: a surviving mutant is
    usually one the benchmark never executes. The blind rate is the checking number.
    """
    document.add_heading("A.1 Outcome decomposition", level=3)
    simple_table(
        document,
        ["outcome", "mutants"],
        [
            ["unobservable (nothing the pack runs reaches it)", metrics.get("unobservable")],
            ["observable, caught by a shipped check", metrics.get("caught_by_pack")],
            ["observable, caught by nothing shipped", metrics.get("unchecked")],
            ["blind rate", metrics.get("blind_rate")],
        ],
    )
    document.add_heading("A.2 Killing layer", level=3)
    simple_table(
        document,
        ["layer", "mutants"],
        [[k, v] for k, v in (metrics.get("by_layer") or {}).items()],
    )
    document.add_heading("A.3 By operator", level=3)
    simple_table(
        document,
        ["operator", "mutants", "survived", "survival rate"],
        [
            [name, row["mutants"], row["survived"], row["survival_rate"]]
            for name, row in (metrics.get("by_operator") or {}).items()
        ],
    )


APPENDIX = {
    "a0": appendix_a0,
    "a1": appendix_a1,
    "a2": appendix_a2,
    "a3": appendix_a3,
    "a4": appendix_a4,
    "a5": appendix_a5,
    "a6": appendix_a6,
    "findings": appendix_all,
}


def appendix_generic(document: Document, metrics: dict[str, Any]) -> None:
    """Fallback appendix: the arm's own top-level scalar readouts.

    A2 and A3 write bespoke report structures rather than the shared metrics schema,
    so their narrative already carries the full tables and this only pins the summary.
    """
    flat = {k: v for k, v in metrics.items() if isinstance(v, (str, int, float, bool)) or v is None}
    if flat:
        document.add_heading("A.1 Summary readouts", level=3)
        simple_table(document, ["key", "value"], [[k, v] for k, v in sorted(flat.items())])
    document.add_paragraph(
        "The complete machine-readable result set for this arm is the metrics.json listed "
        "in the reproduction section. Every table in the body above is read from it."
    )


# A2_rerun shares A2's bespoke schema, so the generic summary appendix is the right one.
# Registered here rather than in the literal above because `appendix_generic` is defined
# after it.
APPENDIX["a2_rerun"] = appendix_generic


# --------------------------------------------------------------------------------
# document assembly
# --------------------------------------------------------------------------------


def digest(path: Path) -> str:
    return "sha256:" + hashlib.sha256(path.read_bytes()).hexdigest()[:16] if path.exists() else "(absent)"


def cover(document: Document, arm: str, metrics: dict[str, Any] | None) -> None:
    label, title, subtitle = ARMS[arm]
    heading = document.add_heading("", level=0)
    add_inline(heading, f"BFCL Oracle-Pack Ablation — {label}")
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(subtitle)
    run.italic = True
    run.font.color.rgb = MUTED

    facts = [
        ["Oracle pack", "banking_vn (9 tools, 17 templates, 6 categories, 9 turn policies)"],
        ["Pipeline", "runtime/benchmark_families/bfcl, unmodified"],
        ["Metric contract", (metrics or {}).get("metrics_version", "1.0")],
    ]
    if arm == "a2_rerun":
        facts.append(["Model", "openai/gpt-oss-120b — but 0 calls made; the run is served entirely from cache"])
    elif arm == "a6":
        facts.append(["Model", "none — this arm is fully deterministic"])
    elif arm == "a5":
        # A5 is the one arm where the model is the subject rather than an authoring tool,
        # so the cover has to say which side of the experiment it sits on.
        facts.append(["Target model", "openai/gpt-oss-120b, local vLLM, temperature 0, /v1/responses, all calls disk-cached"])
    elif arm in {"a2", "a3", "a4", "findings"}:
        facts.append(["Model", "openai/gpt-oss-120b, local vLLM, temperature 0, seed 0, all calls disk-cached"])
    else:
        facts.append(["Model", "none — this arm is fully deterministic"])
    facts.append(["Repository", "Nemotron @ hoannguyen/BFCL"])
    add_table(document, [["field", "value"]] + facts, [WD_ALIGN_PARAGRAPH.LEFT] * 2)


def overview(document: Document, metrics: dict[str, Any]) -> bool:
    """One row per metric: what it measures, and what this arm measured.

    A reader who opens the document cold needs the shape of the result before the
    definitions and the narrative, otherwise the first thing they meet is a five-page
    metric contract. Every figure is computed from the same JSON the appendix reads, so
    the summary cannot disagree with the tables it summarises.

    Returns False when the arm does not use the shared metric schema, so the caller can
    skip the section rather than print an empty one.
    """
    required = ("authoring_friction", "distribution", "coverage", "surface", "funnel")
    if not all(key in metrics for key in required):
        return False

    friction = metrics["authoring_friction"]
    distribution = metrics["distribution"]
    coverage = metrics["coverage"]
    surface = metrics["surface"]["overall"]
    funnel = metrics["funnel"]

    shares = distribution["policy_task_share"]
    top_policy = max(shares, key=shares.get) if shares else "-"
    dropped = funnel["expanded"] - funnel["published"]

    rows = [
        [
            "Authoring friction",
            "Lines a person must write to stand up the pack",
            f"{friction['loc_total']} lines; {friction['template_count']} templates, "
            f"{friction['category_count']} categories, {friction['policy_count']} policies",
        ],
        [
            "Joint (category x policy) distribution",
            "Whether coverage of conversation shapes is a target or an accident",
            f"{distribution['cells_populated']}/{distribution['cells_total']} cells populated; "
            f"{top_policy} {shares.get(top_policy, 0):.1%}",
        ],
        [
            "Slot and fixture coverage",
            "How much of the domain is exercised, including entities never touched",
            f"{coverage['fixture_entities_bound']}/{coverage['fixture_entities_total']} entities bound; "
            f"{len(coverage['tools_declared']) - len(coverage['tools_never_called'])}"
            f"/{len(coverage['tools_declared'])} tools called",
        ],
        [
            "Utterance diversity",
            "How many genuinely different sentences, as opposed to how many rows",
            f"{surface['tasks']} tasks -> {surface['distinct_masked']} distinct slot-masked "
            f"({surface['surfaces_per_template']} per template)",
        ],
        [
            "Publish funnel",
            "Where tasks drop out between expansion and publication, and why",
            f"{funnel['expanded']} -> {funnel['published']}, "
            f"{'nothing dropped' if dropped == 0 else f'{dropped} dropped'}; "
            f"publish {funnel['publish_rate']:.0%}, gold {funnel['gold_rate']:.0%}",
        ],
    ]
    document.add_heading("At a glance", level=1)
    document.add_paragraph(
        "Each metric is defined in Part 1 and reported in full in Appendix A. The figures "
        "below are computed from the same result set the appendix reads."
    )
    add_table(
        document,
        [["metric", "what it measures", "result"]] + rows,
        [WD_ALIGN_PARAGRAPH.LEFT] * 3,
    )
    return True


def reproduction(document: Document, arm: str) -> None:
    document.add_heading("Reproduction", level=2)
    commands = {
        "a0": "PYTHONPATH=src python3 bfcl_ablation/run_a0.py\nPYTHONPATH=src python3 bfcl_ablation/sweep_budget.py 6 12 24",
        "a1": "PYTHONPATH=src python3 bfcl_ablation/run_a1.py",
        "a2": "PYTHONPATH=src python3 bfcl_ablation/run_a2.py",
        "a3": "PYTHONPATH=src python3 bfcl_ablation/run_a3.py",
        "a4": "PYTHONPATH=src python3 bfcl_ablation/run_a4.py",
        "a5": "PYTHONPATH=src python3 bfcl_ablation/run_a5.py",
        "a6": "PYTHONPATH=src python3 bfcl_ablation/run_a6.py",
        "a2_rerun": "PYTHONPATH=src python3 bfcl_ablation/results/A2_rerun/driver.py\nPYTHONPATH=src python3 bfcl_ablation/results/A2_rerun/driver2.py",
        "findings": "PYTHONPATH=src python3 bfcl_ablation/run_a0.py   # then a1, a2, a3, a4, a5, a6",
    }
    document.add_paragraph("Run from the repository root:")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    run = paragraph.add_run(commands[arm])
    run.font.name = MONO
    run.font.size = Pt(9)

    document.add_paragraph(
        "No installation is required: the BFCL family imports with pyarrow, pydantic, "
        "pyyaml and rich alone. run_a1.py exits non-zero when equivalence fails, so it "
        "doubles as a regression test on auto-derivation."
    )

    directory = common.RESULTS / ARMS[arm][0] if arm != "findings" else common.RESULTS
    document.add_heading("Artifacts", level=3)
    rows = []
    if directory.exists():
        for path in sorted(directory.rglob("*")):
            if path.is_file():
                rows.append([str(path.relative_to(common.REPO_ROOT)), f"{path.stat().st_size:,} B", digest(path)])
    metrics_doc = common.RESULTS / "METRICS.md"
    rows.append([str(metrics_doc.relative_to(common.REPO_ROOT)), f"{metrics_doc.stat().st_size:,} B", digest(metrics_doc)])
    simple_table(document, ["path", "size", "digest"], rows)


def build(arm: str) -> Path:
    source = EXPERIMENTS / f"{arm}.md"
    if not source.exists():
        raise FileNotFoundError(source)

    label = ARMS[arm][0]
    metrics = None
    if arm != "findings":
        metrics = common.read_json(common.RESULTS / label / "metrics.json")

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    cover(document, arm, metrics)

    # The summary goes before the definitions: a reader who opens the document cold
    # should meet the shape of the result first, not a five-page metric contract.
    if metrics and overview(document, metrics):
        pass

    # The metric contract travels inside every report. A circulated document whose
    # definitions live in a repo file is a document its reader cannot check: "17
    # distinct masked surfaces" means nothing without the masking rule, and the whole
    # point of pinning definitions is that later arms compare against them.
    contract = common.RESULTS / "METRICS.md"
    if contract.exists():
        document.add_page_break()
        document.add_heading("Part 1 — Method and metric definitions", level=1)
        document.add_paragraph(
            "Shared by every arm of the ablation and reproduced here in full, so this report "
            "stands on its own and needs no other file to be read or checked."
        )
        render_markdown(document, contract.read_text(encoding="utf-8"))

    document.add_page_break()
    document.add_heading("Part 2 — Findings", level=1)
    render_markdown(document, source.read_text(encoding="utf-8"))

    builder = APPENDIX.get(arm)
    if builder is not None:
        document.add_page_break()
        document.add_heading("Appendix A — Full result tables", level=1)
        document.add_paragraph(
            "Every figure below is read from the stored result set at render time rather "
            "than retyped, so this appendix cannot drift from the data it describes."
        )
        builder(document, metrics or {})

    document.add_page_break()
    document.add_heading("Appendix B — Reproduction and lineage", level=1)
    reproduction(document, arm)

    DOCX_OUT.mkdir(parents=True, exist_ok=True)
    target = DOCX_OUT / f"BFCL_Ablation_{label}.docx"
    document.save(target)
    return target


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--arm", choices=sorted(ARMS), help="render one arm (default: all)")
    args = parser.parse_args()

    targets = [args.arm] if args.arm else list(ARMS)
    for arm in targets:
        path = build(arm)
        print(f"wrote {common.rel(path)}  ({path.stat().st_size:,} B)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
